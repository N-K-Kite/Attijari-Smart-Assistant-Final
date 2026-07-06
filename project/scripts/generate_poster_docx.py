import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set margins (padding) of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clear_cell_borders(cell):
    """Remove borders from a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_card_header(cell, icon, title, bg_color):
    """Add a header to a card cell."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    
    # Border-like paragraph style or colored strip
    run_icon = p.add_run(f"{icon}  ")
    run_icon.font.name = 'Arial'
    run_icon.font.size = Pt(14)
    
    run_title = p.add_run(title)
    run_title.font.name = 'Outfit'
    run_title.font.size = Pt(12)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def create_poster_docx():
    doc = Document()
    
    # ── Page Setup: A3 Portrait ──
    # Standard poster format in Word is best at A3
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A3 width in inches
    section.page_height = Inches(16.54) # A3 height in inches
    
    # Set narrow margins (0.4 inch)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    
    # Colors
    c_dark = "1A1A2E"
    c_gold = "C8A04A"
    c_gold_bg = "FAF6EE"
    c_red_bg = "FDF5F6"
    c_blue_bg = "F5F7FA"
    c_purple_bg = "F9F5FD"
    c_green_bg = "F5FDF7"
    
    # ── HEADER TABLE (1 row, 1 col) ──
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    set_cell_background(header_cell, c_dark)
    set_cell_margins(header_cell, top=200, bottom=200, left=300, right=300)
    clear_cell_borders(header_cell)
    
    # Header Content
    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    badge_run = hp.add_run("PROJET DE FIN D'ÉTUDES 2025–2026   ·   SUJET 21   ·   GÉNIE LOGICIEL\n\n")
    badge_run.font.name = 'Arial'
    badge_run.font.size = Pt(9.5)
    badge_run.bold = True
    badge_run.font.color.rgb = RGBColor(0xC8, 0xA0, 0x4A)
    
    title_run = hp.add_run("Système Intelligent de Détection d'Anomalies IT par IA & RPA\n")
    title_run.font.name = 'Outfit'
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    sub_run = hp.add_run("\nPlateforme intégrée combinant Deep Learning (LSTM), Machine Learning (KNN), NLP et automatisation RPA (UiPath)\npour la gestion proactive des incidents critiques à Attijari bank.\n\n")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xCC)
    
    meta_run = hp.add_run("FastAPI · Angular 19 · UiPath   |   1507 tickets réels   |   LSTM Accuracy : 87%   |   AUC : 0.91")
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(10)
    meta_run.bold = True
    meta_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph() # Spacing
    
    # ── MAIN CONTENT GRID (3 Columns) ──
    grid_table = doc.add_table(rows=1, cols=3)
    grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_cell_borders(grid_table.cell(0, 0))
    clear_cell_borders(grid_table.cell(0, 1))
    clear_cell_borders(grid_table.cell(0, 2))
    
    # Col widths (approximate ratio: 1 : 1.25 : 1)
    # Total available width = 11.69 - 0.8 = 10.89 inches
    col_widths = [Inches(3.2), Inches(4.39), Inches(3.3)]
    for i, col in enumerate(grid_table.columns):
        col.width = col_widths[i]
        
    c1 = grid_table.cell(0, 0)
    c2 = grid_table.cell(0, 1)
    c3 = grid_table.cell(0, 2)
    
    # Set padding for grid cells
    for cell in [c1, c2, c3]:
        set_cell_margins(cell, top=0, bottom=0, left=100, right=100)
    
    # ── COLUMN 1: Contexte + Problématique ──
    # Contexte Card
    contexte_table = c1.add_table(rows=1, cols=1)
    ctx_cell = contexte_table.cell(0, 0)
    set_cell_background(ctx_cell, c_gold_bg)
    set_cell_margins(ctx_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(ctx_cell, "🏦", "Contexte du Projet", c_gold)
    
    cp = ctx_cell.add_paragraph()
    cp.paragraph_format.space_after = Pt(6)
    cp.paragraph_format.line_spacing = 1.15
    run = cp.add_run("Attijari bank fait face à un volume croissant d'incidents IT critiques impactant la continuité de ses services bancaires.\n\n"
                     "Ce projet conçoit un système intelligent de détection et résolution automatique des anomalies à partir de 1 507 tickets réels (Février-Mars 2026).\n\n"
                     "La solution couple l'IA (LSTM, NLP) avec le RPA (UiPath) pour un traitement proactif et automatisé en temps réel.")
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    c1.add_paragraph() # Spacing
    
    # Problématique Card
    prob_table = c1.add_table(rows=1, cols=1)
    prob_cell = prob_table.cell(0, 0)
    set_cell_background(prob_cell, c_red_bg)
    set_cell_margins(prob_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(prob_cell, "⚡", "Problématique", c_red_bg)
    
    probs = [
        "Traitement manuel et réactif lent",
        "Absence de détection précoce d'anomalies",
        "Routage manuel des tickets complexe",
        "Temps de résolution (MTTR) élevé",
        "Besoin de suivi des KPI en temps réel"
    ]
    for p_text in probs:
        p_item = prob_cell.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(4)
        run = p_item.add_run(p_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    # ── COLUMN 2: Architecture (Center) ──
    arch_table = c2.add_table(rows=1, cols=1)
    arch_cell = arch_table.cell(0, 0)
    set_cell_background(arch_cell, c_blue_bg)
    set_cell_margins(arch_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(arch_cell, "🏗️", "Architecture du Système", c_blue_bg)
    
    layers = [
        ("FRONTEND", "🅰️ Angular 19  ·  📊 Chart.js  ·  🔐 JWT Auth"),
        ("BACKEND", "⚡ FastAPI  ·  🔑 bcrypt / AES  ·  📧 SMTP Notif"),
        ("IA / ML", "🧠 LSTM (TensorFlow)  ·  🎯 KNN (Similarité)  ·  📖 spaCy"),
        ("RPA", "🤖 UiPath Robot  ·  🔍 CheckAlerte  ·  ✅ Clôture Auto"),
        ("STOCKAGE", "🐘 PostgreSQL  ·  ⚡ Redis Cache  ·  📁 1507 CSV")
    ]
    
    ap = arch_cell.add_paragraph()
    ap.paragraph_format.space_after = Pt(6)
    
    for layer_name, layer_tech in layers:
        l_run_name = ap.add_run(f"■ {layer_name}\n")
        l_run_name.font.name = 'Outfit'
        l_run_name.font.size = Pt(9.5)
        l_run_name.bold = True
        l_run_name.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
        
        l_run_tech = ap.add_run(f"  {layer_tech}\n\n")
        l_run_tech.font.name = 'Arial'
        l_run_tech.font.size = Pt(9)
        l_run_tech.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        
    # Compact Flow Line
    fp = arch_cell.add_paragraph()
    fp.paragraph_format.space_before = Pt(8)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow_run = fp.add_run("FLUX D'INCIDENT\n"
                          "Ticket Créé ➔ NLP spaCy ➔ LSTM Risque ➔ RPA UiPath ➔ Confirmation")
    flow_run.font.name = 'Outfit'
    flow_run.font.size = Pt(9)
    flow_run.bold = True
    flow_run.font.color.rgb = RGBColor(0xC8, 0xA0, 0x4A)
    
    # ── COLUMN 3: Solutions + Technologies ──
    # Solutions Card
    sol_table = c3.add_table(rows=1, cols=1)
    sol_cell = sol_table.cell(0, 0)
    set_cell_background(sol_cell, c_gold_bg)
    set_cell_margins(sol_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(sol_cell, "💡", "Solutions Proposées", c_gold_bg)
    
    solutions = [
        ("1. Détection LSTM", "Modèle prédictif estimant le niveau de risque des anomalies."),
        ("2. Recommandation KNN", "Trouve des solutions historiques similaires (NLP + TF-IDF)."),
        ("3. Robot RPA UiPath", "Vérifie l'API, notifie par mail, et clôture automatiquement."),
        ("4. Dashboard Angular", "Visualisation en temps réel des KPI, graphes et tickets.")
    ]
    for s_title, s_desc in solutions:
        sp = sol_cell.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
        run_t = sp.add_run(f"• {s_title}\n")
        run_t.bold = True
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(9.5)
        run_t.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        
        run_d = sp.add_run(f"  {s_desc}")
        run_d.font.name = 'Arial'
        run_d.font.size = Pt(9)
        run_d.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        
    c3.add_paragraph() # Spacing
    
    # Technologies Card
    tech_table = c3.add_table(rows=1, cols=1)
    tech_cell = tech_table.cell(0, 0)
    set_cell_background(tech_cell, c_purple_bg)
    set_cell_margins(tech_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(tech_cell, "🛠️", "Technologies Clés", c_purple_bg)
    
    tp = tech_cell.add_paragraph()
    tp.paragraph_format.space_after = Pt(2)
    
    techs = [
        ("IA & NLP", "TensorFlow, scikit-learn, spaCy"),
        ("API & DB", "FastAPI, PostgreSQL, Redis"),
        ("Frontend", "Angular 19, TypeScript, Chart.js"),
        ("RPA", "UiPath Studio, Docker")
    ]
    for t_cat, t_list in techs:
        run_cat = tp.add_run(f"▪ {t_cat}: ")
        run_cat.bold = True
        run_cat.font.name = 'Arial'
        run_cat.font.size = Pt(9.5)
        
        run_list = tp.add_run(f"{t_list}\n")
        run_list.font.name = 'Arial'
        run_list.font.size = Pt(9)
        run_list.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        
    doc.add_paragraph() # Spacing
    
    # ── BOTTOM ROW (2 Columns: Résultats & Bénéfices) ──
    bottom_table = doc.add_table(rows=1, cols=2)
    bottom_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_cell_borders(bottom_table.cell(0, 0))
    clear_cell_borders(bottom_table.cell(0, 1))
    
    # Bottom widths (approx ratio: 1.25 : 1)
    # Total available width = 10.89 inches
    bottom_table.columns[0].width = Inches(6.0)
    bottom_table.columns[1].width = Inches(4.89)
    
    bc1 = bottom_table.cell(0, 0)
    bc2 = bottom_table.cell(0, 1)
    
    set_cell_margins(bc1, top=0, bottom=0, left=100, right=100)
    set_cell_margins(bc2, top=0, bottom=0, left=100, right=100)
    
    # Résultats Card
    res_card_table = bc1.add_table(rows=1, cols=1)
    res_cell = res_card_table.cell(0, 0)
    set_cell_background(res_cell, c_green_bg)
    set_cell_margins(res_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(res_cell, "📊", "Résultats & Métriques", c_green_bg)
    
    results = [
        ("87% Précision", "Performance du modèle LSTM sur la détection d'anomalies."),
        ("0.91 AUC-ROC", "Capacité de discrimination excellente du classifieur."),
        ("1507 Tickets", "Tickets réels de production analysés et indexés."),
        ("< 6 secondes", "Temps d'exécution moyen du robot RPA pour la résolution complète.")
    ]
    for r_metric, r_desc in results:
        rp = res_cell.add_paragraph()
        rp.paragraph_format.space_after = Pt(3)
        
        m_run = rp.add_run(f"✓ {r_metric} — ")
        m_run.bold = True
        m_run.font.name = 'Arial'
        m_run.font.size = Pt(9.5)
        m_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        
        d_run = rp.add_run(r_desc)
        d_run.font.name = 'Arial'
        d_run.font.size = Pt(9)
        d_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        
    # Bénéfices Card
    ben_card_table = bc2.add_table(rows=1, cols=1)
    ben_cell = ben_card_table.cell(0, 0)
    set_cell_background(ben_cell, c_green_bg)
    set_cell_margins(ben_cell, top=150, bottom=150, left=150, right=150)
    add_card_header(ben_cell, "✅", "Bénéfices Clés", c_green_bg)
    
    benefits = [
        "Détection proactive avant impact métier",
        "MTTR réduit significativement par RPA",
        "Apprentissage continu du modèle LSTM",
        "Approche Human-in-the-Loop préservée",
        "Traçabilité complète via un Audit Trail",
        "Architecture microservices scalable"
    ]
    for b_text in benefits:
        bp = ben_cell.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        
        b_run = bp.add_run(b_text)
        b_run.font.name = 'Arial'
        b_run.font.size = Pt(9.5)
        b_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    doc.add_paragraph() # Spacing
    
    # ── FOOTER TABLE ──
    footer_table = doc.add_table(rows=1, cols=1)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_cell = footer_table.cell(0, 0)
    set_cell_background(footer_cell, c_dark)
    set_cell_margins(footer_cell, top=100, bottom=100, left=300, right=300)
    clear_cell_borders(footer_cell)
    
    fp = footer_cell.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    foot_run = fp.add_run("Réalisé par : Meriam   |   Encadré par : Attijari bank   |   Établissement : SESAME — Génie Logiciel   |   Année : 2025–2026")
    foot_run.font.name = 'Outfit'
    foot_run.font.size = Pt(9.5)
    foot_run.bold = True
    foot_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Save document
    output_path = r"c:\Users\21655\Desktop\attijari-pfe\poster.docx"
    doc.save(output_path)
    print(f"Poster DOCX generated successfully at: {output_path}")

if __name__ == "__main__":
    create_poster_docx()
